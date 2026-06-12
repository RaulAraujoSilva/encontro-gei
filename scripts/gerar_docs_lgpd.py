# -*- coding: utf-8 -*-
"""
Gera os 3 documentos Word do pacote LGPD do 1º Encontro GEI (tasks 9.5.1–9.5.4):
  1. docs/lgpd/termo-consentimento-imagem-voz.docx
  2. docs/lgpd/aviso-gravacao-transmissao-A4.docx
  3. docs/lgpd/politica-guarda-retencao-dados.docx

Placeholders a validar juridicamente ficam destacados em amarelo.
Uso: python scripts/gerar_docs_lgpd.py
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x09, 0x11, 0x36)
GREEN = RGBColor(0x4E, 0x8A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)

BASE = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "lgpd")

EVENTO = "1º Encontro de Governança, Estratégia e Inovação Governamental"
DATAS = "08, 09 e 10 de julho de 2026 — Rio de Janeiro / Niterói, RJ"
CONTROLADORES = (
    "Agência Reguladora de Energia e Saneamento Básico do Estado do Rio de Janeiro — AGENERSA "
    "e Universidade Federal Fluminense — UFF (PPGEP/UFF), na condição de controladoras conjuntas"
)
AGENERSA_ID = ("AGENERSA — CNPJ 07.694.194/0001-11, Av. Treze de Maio, 23 — Centro, "
               "Rio de Janeiro/RJ, CEP 20031-902")
UFF_ID = ("UFF — autarquia federal, CNPJ 28.523.215/0001-06, Rua Miguel de Frias, 9 — Icaraí, "
          "Niterói/RJ, CEP 24220-900")
EMAIL = "contato@encontrogeig.org"
URL_POLITICA = "encontrogeig.org/privacidade"

CHECKBOX_EVEN3 = (
    "Li e concordo com a Política de Privacidade do Encontro GEI (encontrogeig.org/privacidade). "
    "Estou ciente de que o evento será gravado, fotografado e transmitido ao vivo, e autorizo, "
    "gratuitamente, o uso da minha imagem, voz e nome nos registros e materiais institucionais e "
    "científicos do evento, sem finalidade comercial, nos termos da Lei nº 13.709/2018 (LGPD) e do "
    "art. 20 do Código Civil."
)

PLACEHOLDER = re.compile(r"(\[[^\]]*A (?:VALIDAR|PREENCHER)[^\]]*\])")


def add_runs(par, text, size=11, bold=False, color=None, italic=False):
    """Adiciona texto ao parágrafo destacando placeholders [ ... A VALIDAR/PREENCHER ...] em amarelo."""
    for part in PLACEHOLDER.split(text):
        if not part:
            continue
        run = par.add_run(part)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if PLACEHOLDER.fullmatch(part):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.bold = True
    return par


def p(doc, text, size=11, bold=False, color=None, align=None, space_after=8, italic=False):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        par.alignment = align
    add_runs(par, text, size=size, bold=bold, color=color, italic=italic)
    return par


def heading(doc, text, size=14, color=NAVY, space_before=14):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return par


def title_block(doc, title, subtitle):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(EVENTO)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = GREEN
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(16)
    run = par.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Arial")


def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def make_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        par = cell.paragraphs[0]
        run = par.add_run(header)
        run.font.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "091136")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], value, size=font_size)
    if widths:
        for i, width in enumerate(widths):
            for row_obj in table.rows:
                row_obj.cells[i].width = Cm(width)
    return table


# ----------------------------------------------------------------------------
# Documento 1 — Termo de Consentimento de Uso de Imagem e Voz
# ----------------------------------------------------------------------------
def doc_termo():
    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title_block(
        doc,
        "TERMO DE AUTORIZAÇÃO DE USO DE IMAGEM, VOZ E NOME",
        f"{DATAS} · Documento sujeito a validação jurídica — versão 1.0 de 12/06/2026",
    )

    heading(doc, "1. Identificação")
    p(doc, f"Evento: {EVENTO} (“Evento”), realizado nos dias {DATAS}, em formato híbrido "
           "(presencial e online, com transmissão ao vivo).")
    p(doc, f"Organização e controladoras dos dados pessoais: {CONTROLADORES}. "
           f"{AGENERSA_ID}. {UFF_ID}.")

    heading(doc, "2. Considerandos")
    p(doc, "Considerando que o Evento é gratuito, de natureza institucional, acadêmica e científica; "
           "que será integralmente fotografado, gravado e transmitido ao vivo pela internet (YouTube); "
           "e que seus registros integram o acervo institucional e de divulgação científica do Encontro, "
           "o presente Termo formaliza a autorização de uso de imagem, voz e nome do(a) participante.")

    heading(doc, "3. Objeto")
    p(doc, "O(A) signatário(a) AUTORIZA, de forma livre, informada e inequívoca, a captação e o uso de sua "
           "imagem, voz e nome em fotografias, gravações audiovisuais, transmissão ao vivo e materiais "
           "institucionais e científicos do Evento — incluindo vídeos de apresentação de trabalhos — "
           "SEM QUALQUER FINALIDADE COMERCIAL.")

    heading(doc, "4. Abrangência")
    p(doc, "A autorização abrange a divulgação nos canais oficiais do Evento e das instituições realizadoras: "
           "transmissão e acervo no YouTube, site oficial (encontrogeig.org), redes sociais institucionais, "
           "anais e publicações científicas (incluindo livro com ISBN) e materiais de registro e prestação de "
           "contas. Vigência: por prazo indeterminado, enquanto perdurar o acervo institucional, em território "
           "nacional e no exterior (em razão da natureza da internet).")

    heading(doc, "5. Gratuidade")
    p(doc, "A presente autorização é concedida a título gratuito, nada sendo devido ao(à) signatário(a), a "
           "qualquer tempo, a título de direitos de imagem, voz, nome ou conexos.")

    heading(doc, "6. Fundamentos legais")
    p(doc, "Este Termo fundamenta-se no art. 7º, inciso I (consentimento), da Lei nº 13.709/2018 (LGPD), e no "
           "art. 20 do Código Civil (Lei nº 10.406/2002). O tratamento dos demais dados pessoais do(a) "
           f"participante observa a Política de Privacidade do Evento, disponível em {URL_POLITICA}.")

    heading(doc, "7. Revogação")
    p(doc, "Nos termos do art. 8º, § 5º, da LGPD, esta autorização pode ser revogada a qualquer tempo, "
           f"mediante manifestação gratuita e facilitada ao canal {EMAIL}. A revogação produz efeitos "
           "prospectivos: não alcança a transmissão ao vivo já realizada nem os usos consolidados ao amparo "
           "deste Termo, e será atendida na medida do tecnicamente possível quanto a registros panorâmicos "
           "de plateia.")

    heading(doc, "8. Assinatura (versão para credenciamento presencial)")
    p(doc, "Nome completo: ______________________________________________________________")
    p(doc, "CPF (opcional): _______________________  Instituição: _______________________________")
    p(doc, "Local e data: ________________________________________  ____ / ____ / 2026")
    p(doc, "Assinatura: __________________________________________________________________")

    doc.add_page_break()
    heading(doc, "ANEXO — Versões curtas derivadas deste Termo", size=14)

    heading(doc, "A.1 Texto do campo de aceite (checkbox) no formulário de inscrição Even3", size=11, color=GREEN)
    p(doc, f"“{CHECKBOX_EVEN3}”", italic=True)

    heading(doc, "A.2 Texto de balcão (credenciamento)", size=11, color=GREEN)
    p(doc, "“Este evento será gravado, fotografado e transmitido ao vivo. Ao se credenciar, você declara "
           "ciência da captação de imagem e voz nos ambientes do evento, conforme a Política de Privacidade "
           f"({URL_POLITICA}). Dúvidas ou oposição: {EMAIL}.”", italic=True)

    heading(doc, "A.3 Quando usar o termo completo (com assinatura)", size=11, color=GREEN)
    p(doc, "O termo completo com assinatura destina-se a palestrantes, painelistas, moderadores e autores que "
           "apresentam trabalhos (exposição individualizada e permanente). Para o público em geral, bastam o "
           "aceite na inscrição (A.1), o aviso de balcão (A.2) e a sinalização ostensiva dos ambientes.")

    path = os.path.join(BASE, "termo-consentimento-imagem-voz.docx")
    doc.save(path)
    return path


# ----------------------------------------------------------------------------
# Documento 2 — Aviso de Gravação/Transmissão (cartaz A4 + textos online)
# ----------------------------------------------------------------------------
def doc_aviso():
    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Página 1 — cartaz
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(40)
    run = par.add_run(EVENTO)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = GREEN

    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(20)
    par.paragraph_format.space_after = Pt(20)
    run = par.add_run("ESTE EVENTO ESTÁ SENDO\nGRAVADO, FOTOGRAFADO E\nTRANSMITIDO AO VIVO")
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p(doc, "As imagens e os áudios captados destinam-se ao registro institucional, à transmissão pelo YouTube "
        "e ao acervo de divulgação científica do Encontro, sem finalidade comercial.",
        size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    p(doc, "Ao permanecer neste ambiente, você declara ciência da captação de sua imagem e voz.",
        size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    p(doc, f"Saiba mais ou manifeste oposição: {URL_POLITICA} · {EMAIL}",
        size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=8)
    p(doc, "[ ESPAÇO RESERVADO PARA QR CODE DA POLÍTICA DE PRIVACIDADE — A PREENCHER NA DIAGRAMAÇÃO ]",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=30)
    p(doc, "Lei nº 13.709/2018 (LGPD) · art. 20 do Código Civil",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)

    doc.add_page_break()

    # Página 2 — versões para o online
    title_block(doc, "AVISOS PARA A TRANSMISSÃO ONLINE",
                "Textos de apoio — descrição do YouTube, slide de abertura e roteiro do mestre de cerimônias")

    heading(doc, "1. Texto para a descrição do vídeo/transmissão no YouTube", size=12, color=GREEN)
    p(doc, f"“Transmissão oficial do {EVENTO} ({DATAS}). Este evento é gravado e transmitido ao vivo; as "
           "imagens integram o acervo institucional e de divulgação científica do Encontro, sem finalidade "
           f"comercial. Política de Privacidade: https://{URL_POLITICA}”", italic=True)

    heading(doc, "2. Texto para slide/vinheta de abertura da transmissão", size=12, color=GREEN)
    p(doc, "“Este evento está sendo gravado, fotografado e transmitido ao vivo.\n"
           f"Política de Privacidade: {URL_POLITICA}”", italic=True)

    heading(doc, "3. Aviso verbal do mestre de cerimônias (abertura de cada dia)", size=12, color=GREEN)
    p(doc, "“Informamos que este evento está sendo gravado, fotografado e transmitido ao vivo pelo "
           "YouTube. As imagens e os áudios captados integram o acervo institucional e científico do "
           "Encontro, sem finalidade comercial. Mais informações na Política de Privacidade, disponível em "
           f"{URL_POLITICA}.”", italic=True)

    heading(doc, "4. Orientações de aplicação", size=12, color=GREEN)
    p(doc, "• Imprimir a página 1 em A4 (sugere-se também versão A3) e afixar em todos os acessos: "
           "credenciamento, entradas dos auditórios, foyer e pontos de embarque das visitas técnicas.")
    p(doc, "• Exibir o slide de abertura (item 2) antes do início de cada bloco da transmissão.")
    p(doc, "• O aviso verbal (item 3) deve ser lido na abertura de cada dia do evento.")
    p(doc, "• Gerar o QR Code apontando para https://" + URL_POLITICA + " e inseri-lo no espaço reservado do cartaz.")

    path = os.path.join(BASE, "aviso-gravacao-transmissao-A4.docx")
    doc.save(path)
    return path


# ----------------------------------------------------------------------------
# Documento 3 — Política interna de Guarda e Retenção de Dados
# ----------------------------------------------------------------------------
def doc_retencao():
    doc = Document()
    set_base_style(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title_block(
        doc,
        "POLÍTICA INTERNA DE GUARDA E RETENÇÃO DE DADOS PESSOAIS",
        f"{DATAS} · Documento interno da organização — versão 1.0 de 12/06/2026 · sujeito a validação jurídica",
    )

    heading(doc, "1. Objetivo e escopo")
    p(doc, "Definir, para a organização do Evento, as regras de guarda, retenção, acesso e descarte dos dados "
           "pessoais tratados em razão do Evento (inscrição, credenciamento, submissão de trabalhos, presença, "
           "certificação, registros audiovisuais e comunicações), em conformidade com a Lei nº 13.709/2018 "
           "(LGPD), em especial os arts. 6º (princípios), 15 e 16 (término do tratamento), 37 (registro das "
           "operações), 39 (operador), 46 a 49 (segurança) e 48 (incidentes).")

    heading(doc, "2. Papéis e responsabilidades")
    p(doc, f"Controladoras conjuntas: {CONTROLADORES} ({AGENERSA_ID}; {UFF_ID}). "
           "Recomenda-se formalizar acordo entre as controladoras "
           "definindo responsabilidades (transparência, atendimento a titulares, comunicação de incidentes), "
           "conforme o Guia Orientativo da ANPD sobre agentes de tratamento. [MINUTA DO ACORDO — A VALIDAR]")
    p(doc, "Operadora: Even3 S.A. — plataforma de inscrição, submissão e certificação, que trata dados em nome "
           "da organização (art. 39 da LGPD), conforme seus termos de uso e política de privacidade.")
    p(doc, "Plataformas com tratamento próprio: YouTube/Google (transmissão e acervo público).")
    p(doc, "Encarregado (DPO) do Evento: [NOME DO ENCARREGADO — A VALIDAR].")
    p(doc, "Pessoas com acesso ao painel da Even3 e às caixas de e-mail oficiais: [LISTA NOMINAL — A PREENCHER]. "
           "O acesso deve ser limitado ao mínimo necessário e revisto ao final do Evento.")

    heading(doc, "3. Inventário de dados pessoais")
    make_table(
        doc,
        ["Categoria", "Origem", "Finalidade", "Base legal", "Onde fica", "Quem acessa", "Prazo", "Descarte"],
        [
            ["Inscrição (nome, e-mail, instituição, telefone)", "Formulário Even3", "Inscrição, credenciamento, comunicação operacional", "Art. 7º, V", "Even3", "Organização (painel)", "[5 anos — A VALIDAR]", "Eliminação na Even3 + cópias locais"],
            ["Presença (QR Code)", "Credenciamento", "Apuração de 75% p/ certificado", "Art. 7º, V e IX", "Even3", "Organização", "[5 anos — A VALIDAR]", "Junto com a inscrição"],
            ["Submissões (trabalhos, autoria, URL de vídeo)", "Formulário Even3", "Avaliação, programa, anais com ISBN", "Art. 7º, V", "Even3 / anais", "Comissão científica", "Permanente (publicação)", "Não se aplica (acervo)"],
            ["Imagem, voz e nome (registros do evento)", "Captação no evento", "Registro institucional, transmissão, divulgação científica", "Art. 7º, IX e I; CC art. 20", "YouTube, site, redes", "Organização / público", "Permanente (acervo)", "Não se aplica (acervo)"],
            ["E-mails operacionais", "Caixas @encontrogeig.org", "Atendimento e histórico", "Art. 7º, V e IX", "E-mail institucional", "Organização", "[2 anos — A VALIDAR]", "Exclusão das caixas"],
            ["Acessibilidade (se informada)", "Participante", "Atendimento da necessidade", "Art. 7º, I / art. 11, II, a", "Even3 / e-mail", "Logística (restrito)", "Até o fim do evento", "Exclusão imediata"],
        ],
        font_size=8,
    )

    heading(doc, "4. Prazos de retenção")
    p(doc, "Os prazos da tabela acima contam-se do encerramento do Evento. Hipóteses de conservação após o "
           "término do tratamento fundamentam-se no art. 16 da LGPD: inciso I (cumprimento de obrigação legal "
           "ou regulatória — comprovação de certificação e prestação de contas) e inciso II (estudo por órgão "
           "de pesquisa — anais e acervo científico). Prazos marcados como [A VALIDAR] devem ser confirmados "
           "com as assessorias jurídicas da AGENERSA e da UFF.")

    heading(doc, "5. Procedimentos de descarte")
    p(doc, "• Ao término do prazo de retenção: exportar da Even3 apenas o necessário à prestação de contas, "
           "eliminar os registros na plataforma e excluir cópias locais (planilhas, e-mails, drives).")
    p(doc, "• Registrar a eliminação (data, responsável, categoria de dados) em ata simples, arquivada com a "
           "documentação do Evento.")
    p(doc, "• Dados de acessibilidade: excluir imediatamente após o encerramento do Evento.")

    heading(doc, "6. Atendimento aos direitos dos titulares (arts. 18 e 19)")
    p(doc, f"Canal único: {EMAIL}. Responsável pelo fluxo: [RESPONSÁVEL — A PREENCHER]. Registrar cada "
           "solicitação (data, titular, pedido, resposta). Responder em prazo razoável, observando o art. 19 "
           "da LGPD (declaração simplificada imediata ou declaração completa em até 15 dias, no caso de "
           "confirmação de tratamento e acesso). Solicitações que envolvam a Even3 devem ser encaminhadas à "
           "plataforma e acompanhadas até a conclusão.")

    heading(doc, "7. Incidentes de segurança (art. 48 e Resolução CD/ANPD nº 15/2024)")
    p(doc, "Em caso de incidente de segurança com dados pessoais que possa acarretar risco ou dano relevante "
           "aos titulares: (i) registrar o incidente e acionar imediatamente o encarregado; (ii) avaliar o "
           "risco; (iii) sendo relevante, comunicar a ANPD e os titulares afetados em até 3 (três) dias úteis "
           "contados do conhecimento do incidente, pelo formulário eletrônico da ANPD; (iv) complementar "
           "informações em até 20 dias úteis, se necessário; (v) documentar causas, efeitos e medidas adotadas.")

    heading(doc, "8. Vigência e revisão")
    p(doc, "Esta política vigora a partir de sua aprovação pelas controladoras e deve ser revista após o "
           "encerramento do Evento (ou imediatamente, em caso de mudança legislativa ou operacional relevante). "
           "Aprovação: AGENERSA ____/____/2026 · UFF ____/____/2026. [ASSINATURAS — A VALIDAR]")

    path = os.path.join(BASE, "politica-guarda-retencao-dados.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    os.makedirs(BASE, exist_ok=True)
    for fn in (doc_termo, doc_aviso, doc_retencao):
        print("OK:", fn())
